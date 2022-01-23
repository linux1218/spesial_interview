import time
import datetime
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
from docx import Document



global browser
global novel_url_list
global total_info
global stop_subject

def init_webdriver():
    global browser
    # options = webdriver.ChromeOptions()
    # options.headless = True
    # options.add_argument("window-size=1920x1080")
    # options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.110 Safari/537.36")
    # browser = webdriver.Chrome(options=options)
    # browser.implicitly_wait(1)
    browser = webdriver.Chrome()


def clear_webdriver():
    global browser
    browser.quit()


def get_novel_url():
    global browser
    global novel_url_list
    global stop_subject

    novel_url_list=[]


    starturl = "https://novel.naver.com/challenge/list?novelId=1042098"
    
    #제일 마지막화 제목
    browser.get(starturl) # start url 로 이동
    soup = BeautifulSoup(browser.page_source, "lxml")
    subject_list = soup.find("ul", attrs={"class":"list_type2 v3 league_num NE=a:lst"}).find_all("li",attrs={"class":"volumeComment"})
    stop_subject = subject_list[0].find("p", attrs={"class":"subj"}).get_text().strip()
    stop_subject = stop_subject[:stop_subject.find('\n')]


    base_url = "https://novel.naver.com/challenge/list?novelId=1042098&order=Oldest&page="

    #시작 URL로 변경
    page_index=0

    while True:        
        page_index+=1

        url = base_url + str(page_index)
        
        # try:
        browser.get(url) # url 로 이동
        time.sleep(1)

        soup = BeautifulSoup(browser.page_source, "lxml")

        novel_list_elements = soup.find("ul", attrs={"class":"list_type2 v3 league_num NE=a:lst"}).find_all("li", attrs={"class":"volumeComment"})

        for novel_list_element in novel_list_elements:
            novel_url=novel_list_element.find("a", attrs={"class":"list_item NPI=a:list"})["href"]
            novel_url_list.append("https://novel.naver.com"+novel_url)
        
        try:
            nest_page_index=soup.find("div", attrs={"class":"paging NE=a:lst"}).find("a", text=str(page_index+1))
        except Exception as err: # 처음 페이지가 없을때 
            break
        
        if not nest_page_index:
            break



def scrape_novel_info():
    global browser
    global total_info

    total_info=[]

    print("[그와의 은밀한 면접 백업]")

    starturl = "https://novel.naver.com/challenge/list?novelId=1042098"    
    check_index = 0 #루프 횟수 

    while True:
        try:
            browser.get(novel_url_list[check_index])

            time.sleep(1)
            novel_subject = browser.find_element(By.XPATH, "//*[@id='content']/div[1]/div[2]/h2").text
            novel_subject = novel_subject[:novel_subject.find('\n')]

            novel_content = browser.find_element(By.XPATH, "//*[@id='content']/div[1]/div[2]/div/p").text

            # 정보 일괄 취합
            single_info=[novel_subject, novel_content]
            total_info.append(single_info)

            if stop_subject == novel_subject:
                break

            check_index = check_index + 1

        except Exception as err:
            print(err)
            break


def write_ms_word():
    document = Document()
    now = datetime.datetime.now()
    currdate = now.strftime('%Y%m%d_%H시%M분')
    save_file_name=f'C:/PythonWorkSpace/novel_{currdate}.docx'

    for single_info in total_info:
        document.add_heading(single_info[0], level = 2)
        document.add_paragraph(single_info[1])

    document.save(save_file_name)


if __name__ == "__main__":
    try:
        init_webdriver()
        get_novel_url()
        scrape_novel_info()
        write_ms_word()
        clear_webdriver()
    except Exception as err:
        print(err)


